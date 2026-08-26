# SPDX-FileCopyrightText: Copyright 2026 routerx contributors
# SPDX-License-Identifier: Apache-2.0

"""결과보고서 공식 양식(DOCX)에 원고를 채운다.

양식의 표 구조를 유지하고 안내 문구만 실제 내용으로 바꾼다.
본문은 공식 개조식으로 쓰되 소감 및 후기만 서술식으로 둔다.
"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt

SRC = Path(r"D:\opensource\결과보고서양식\2026 오픈소스 개발자대회 결과보고서_접수번호(팀명)"
           r"\2026 오픈소스 개발자대회 결과보고서_접수번호(팀명).docx")
OUT = Path(r"d:\opensource\skt-router\2026 오픈소스 개발자대회 결과보고서_접수번호(트리아지).docx")

TEAM = "트리아지"
MEMBERS = "2명"
REPO = "https://github.com/leelang7/ossp-2026-llm-router-challenge"
# 제출 가이드: 프로젝트 등록 URL에는 대표 저장소 링크 하나만 적는다.
# 심사 기준 시점은 본문 7번 항목에 저장소 태그로 따로 밝힌다.
TAG = "submission-2026"
IMAGE = ("ghcr.io/leelang7/routerx@sha256:"
         "ea01be4aa373f1358450c56105f4f595619b7fa2bd272d418c9bc71f8b75016f")
VIDEO = "[유튜브 URL — 업로드 후 기재]"

BODY = {
    1: "프롬프트 난이도 인지 기반 경량 LLM 라우터",
    2: REPO,
    3: VIDEO,
    4: "프롬프트 내용만으로 후보 모델 하나를 선택하여 주어진 예산 범위에서 답변 품질을 "
       "최대화하는 오픈소스 LLM 라우터. 모델 호출과 네트워크 접근 없이 학습된 계수와 "
       "어휘만으로 문항당 약 5밀리초에 판단.",
    6: "○ 추진 배경\n"
       "  - LLM 서비스 운영비에서 추론 비용이 차지하는 비중 최대\n"
       "  - 실제 트래픽의 상당수는 경량 모델로도 충분한 단순 질의\n"
       "  - 전량을 고성능 모델로 처리하는 관행에 따른 예산 낭비 발생\n"
       "○ 개발 목표\n"
       "  - 프롬프트 난이도를 자동 판별하여 최적 모델에 배정\n"
       "  - 동일 예산 조건에서 평균 답변 품질 향상\n"
       "  - 재현 가능하고 이식 가능한 오픈소스 구현\n"
       "○ 과제 제약\n"
       "  - 입력은 프롬프트 내용과 예산 등급으로 한정\n"
       "  - 모델 호출, 답변 비교, 선택 번복 불가\n"
       "  - 예산 초과 시 해당 등급 0점 처리",
    7: "○ 언어 및 런타임 : Python 3.11, 컨테이너 linux/arm64 (Debian slim)\n"
       "○ 학습 환경 : scikit-learn 1.8.0, SciPy 1.17.1, NumPy 1.26.4\n"
       "○ 실행 의존성 : NumPy 단일 패키지 (학습용 패키지는 이미지 미포함)\n"
       "○ 공식 실행 조건 : CPU 2코어, 메모리 2GiB, 등급당 90초\n"
       "○ 격리 조건 : 네트워크 차단, 읽기 전용 루트, 비특권 사용자(UID 65532)\n"
       "○ 개발 도구 : Git, Docker Buildx, 표준 라이브러리 unittest\n"
       "○ 산출물 규모 : 컨테이너 이미지 73.4MB, 학습 아티팩트 5.5MB\n"
       f"○ 심사 기준 스냅샷 : 저장소 태그 {TAG}\n"
       f"○ 제출 이미지 : {IMAGE}",
    8: "○ 특징 추출 — 프롬프트 내용만 사용\n"
       "  - 직접 계산 특징 36종 : 길이, 한글 비율, 수식·코드 표지, 객관식 표지, 대화 구조\n"
       "  - TF-IDF 단어 1~2그램 60,000차원\n"
       "  - TF-IDF 문자 3~5그램 120,000차원\n"
       "○ 예측 — ridge 회귀 7개 헤드\n"
       "  - 모델별 예측 품질 3종\n"
       "  - 모델별 예측 출력 토큰 3종, 예측 입력 토큰 1종\n"
       "  - 공개 비용 정책을 적용하여 모델별 비용으로 환산\n"
       "○ 배치 선택 정책\n"
       "  - 예측 이득 대비 예측 비용이 큰 순서로 승격\n"
       "  - 동률은 프롬프트 내용 해시로만 정렬 (문항 ID·입력 순서 미사용)\n"
       "  - 등급별 예산 마진, 추론 모델 선택 건수 상한, 문항별 비용 상한 적용\n"
       "○ 출력 — 등급별 선택 결과 JSON (문항당 model_id 1개)",
    9: "□ 비용의 예측 대상 포함\n"
       "  - 라우팅 시점에 토큰 수가 제공되지 않아 비용 자체를 예측 대상에 포함\n"
       "  - 학습 단계에서도 실제 토큰 수 대신 예측값을 사용하여 추론 조건과 일치\n"
       "  - 로그 공간 회귀의 지수 환원 시 발생하는 평균 과소평가를 잔차 분위수로 보정\n"
       "  - 학습 분할의 실제 총비용을 기준으로 모델별 스케일 조정\n"
       "□ 예산 초과의 구조적 차단\n"
       "  - 추론 모델 출력 토큰 분포 : 중앙값 1,570 / 최댓값 130,504\n"
       "  - 단일 문항이 경량 모델 전체 비용의 26%까지 점유 가능\n"
       "  - 예측 비용 상한 및 분위수 회귀 방식 : 미채택\n"
       "    · 사고 유발 문항이 예측 과소·실제 과대 유형이어서 필터를 통과\n"
       "  - 선택 건수 상한 방식 : 채택\n"
       "    · 예측 오차와 무관하게 노출 규모를 제한\n"
       "    · 고비용 구간은 승격 이득 대비 비용도 최저이므로 점수 손실 경미\n"
       "□ 순서 비의존 배치 정책\n"
       "  - 예산은 배치 전체 제약이므로 문항 단위 독립 처리 불가\n"
       "  - 입력 순서에 의존할 경우 과제 규칙 위반에 해당\n"
       "  - 승격 후보를 이득·비용비로 정렬하되 동률은 프롬프트 SHA-256 해시로 판정\n"
       "  - 내용이 동일한 문항은 그룹 단위로 일괄 승격\n"
       "  - 문항 ID 및 입력 순서 변경 후 재실행 시에도 동일 결과 보장\n"
       "□ 학습·실행 경로 분리\n"
       "  - 학습은 scikit-learn으로 수행하고 제출 이미지에는 NumPy만 포함\n"
       "  - TfidfVectorizer 동작을 표준 라이브러리로 재현 (소문자화, 토큰 패턴,\n"
       "    char_wb 패딩, sublinear tf, smooth idf, 블록별 L2 정규화)\n"
       "  - 두 경로의 예측 차이 1.2e-15 (부동소수점 한계 수준)\n"
       "□ 등급별 설정 — 교차검증 기반 결정\n"
       "  - 2,640문항 대상 5-fold 및 8-fold 교차검증 수행\n"
       "  - 양쪽 fold 구성에서 예산 초과 0건인 조합만 채택\n"
       "  - fast : 추론모델 상한 0%, 예산 마진 0.85 (Dev 사용률 82.9%)\n"
       "  - balanced : 추론모델 상한 1%, 예산 마진 0.80 (Dev 사용률 80.1%)\n"
       "  - premium : 추론모델 상한 11%, 예산 마진 0.85 (Dev 사용률 64.7%)\n"
       "  - 문항별 추론모델 비용 상한 5% 병행 적용\n"
       "□ 구동 방법\n"
       "  - 학습 : python train_routerx/train.py --artifact src/routerx/artifact.npz\n"
       "  - 라우팅 : python -m routerx.cli --input INPUT --tier TIER --output OUTPUT\n"
       "  - 자체 점검 : python -m routerx.audit --input INPUT --outcomes OUTCOMES\n"
       "  - 컨테이너 : 공식 자원 한도 조건에서 실행 확인 완료\n"
       "□ 자체 점검 결과 — 공개 Dev 880문항\n"
       "  - ID·순서 불변성 : 400문항 불일치 0건\n"
       "  - 결정성 : 300문항 불일치 0건\n"
       "  - 엣지 케이스 7종 통과 (빈 프롬프트, 초장문, 어휘 외 문자, 이모지, 장문 대화)\n"
       "  - 실행 시간 : 등급당 4.2초 (한도 90초 대비 5%)\n"
       "  - 예산 : 3개 등급 전량 통과, 여유 17~35%\n"
       "  - 단위 시험 21건 통과\n"
       "□ 성능\n"
       "  - 공개 Dev 880문항 : 최종 0.7167\n"
       "  - 교차검증 2,640문항 : 0.6643 (예산 초과 0건)\n"
       "  - 공식 최고 baseline 0.6954 / 전량 경량 모델 0.6193",
    10: "○ 이식성\n"
        "  - 특정 모델군에 비종속하며, 후보 모델의 품질·토큰 기록과 비용 정책만으로\n"
        "    임의의 모델 조합에 재학습 가능\n"
        "  - 학습 스크립트 및 평가 도구를 함께 공개하여 사내 모델 구성 기준 재현 지원\n"
        "○ 운영 적합성\n"
        "  - 네트워크 및 GPU 불필요, CPU 2코어에서 동작하여 사내망·온프레미스 즉시 적용\n"
        "  - 아티팩트 5.5MB로 엣지 환경 탑재 가능\n"
        "○ 기대효과\n"
        "  - 고객센터·검색·공공 AI 서비스 등 대규모 트래픽 환경의 추론 비용 절감\n"
        "  - 동일 예산에서 품질 향상 또는 동일 품질에서 비용 절감\n"
        "  - 라우팅 정책·학습 파이프라인·평가 도구 공개로 관련 연구의 재현 기준점 제공",
    11: "□ 혁신성 및 차별성\n"
        "  - 예산 초과를 예측 정확도가 아닌 구조적 장치로 차단\n"
        "    · 예측 기반 방어의 실패 원인을 실험으로 규명한 후 예측 무관 장치로 대체\n"
        "  - 순서 불변성을 설계 단계에서 보장\n"
        "    · 동률 정렬 키를 프롬프트 내용 해시로 지정하여 감사 항목을 구조적으로 충족\n"
        "  - 학습·실행 경로의 일치를 수치로 증명하고 시험으로 고정 (차이 1.2e-15)\n"
        "  - 제출 전 자체 점검 도구를 함께 제공\n"
        "    · 규칙 준수, 결정성, 엣지 케이스, 실행 시간, 예산 여유를 일괄 확인\n"
        "  - 정보 수준별 도달 가능 상한을 분석하여 개선 여지를 정량화\n"
        "    · 전량 경량 0.605 / 비용 정확 0.683 / 이득 완전정보 0.725 / 완전정보 0.792\n"
        "  - 측정 결과에 근거한 기술 선택 및 설계 회귀 2건\n"
        "    · 부스팅 트리 : 단일 측정 최고점이나 교차검증 예산 미충족으로 미채택\n"
        "    · 문장 임베딩 : 예측 성능 저하(0.425 → 0.372)로 미채택\n"
        "□ 한계점\n"
        "  - 경량·중형 모델 간 품질 차이는 예측 상관 0.01~0.10으로 사실상 예측 불가\n"
        "    · 점수가 생성 2~4회 평균값이어서 확률적 변동이 지배적\n"
        "    · 부스팅 트리 및 다국어 임베딩 적용으로도 한계 미극복\n"
        "  - 안전 확보를 위해 예산의 17~35%를 미사용 상태로 유지\n"
        "  - 컨테이너 이미지 73MB (NumPy 의존)\n"
        "□ 향후 발전 로드맵\n"
        "  - 1단계 : 출력 길이 예측 정확도 개선을 통한 예산 활용률 제고\n"
        "  - 2단계 : 실시간 서빙용 온라인 예산 컨트롤러 구현\n"
        "  - 3단계 : 타 모델군 및 비용 정책으로의 이식 사례 확보\n"
        "  - 4단계 : 생성 반복 횟수가 확대된 자료 확보 시 품질 예측 재검증\n"
        "\n"
        "□ 소감 및 후기\n"
        "본 과제는 성능 향상보다 안정성 확보에 더 많은 시간이 소요되었다. 예산을 최대한 "
        "활용하는 정책은 공개 자료에서 우수한 결과를 보였으나 데이터 분할을 달리하면 한도를 "
        "초과하였고, 그때마다 해당 등급의 점수가 전부 무효화되었다. 평균값을 조정하는 마진 "
        "설정만으로는 분포의 두꺼운 꼬리를 통제할 수 없다는 사실을 여러 차례의 실패를 거쳐 "
        "확인한 뒤에야, 위험을 예측하려 시도하는 대신 노출 규모 자체를 제한하는 방향으로 "
        "설계를 전환할 수 있었다.\n"
        "\n"
        "학습 코드와 실행 코드가 분리된 지점에서 예측이 어긋난 결함 또한 기억에 남는다. "
        "공식 검증 절차는 통과하였으나 실제로는 세 등급 모두 예산을 초과하는 상태였으며, "
        "자체 점검 도구를 별도로 마련하지 않았다면 제출 이후에야 발견하였을 것이다. 검증 "
        "절차를 코드로 남기는 작업이 기능 구현 못지않게 중요하다는 점을 체감하였다.\n"
        "\n"
        "측정 결과가 예상과 어긋날 때 이를 수용하는 판단도 쉽지 않았다. 부스팅 트리와 문장 "
        "임베딩은 모두 개선이 기대되는 접근이었으나 교차검증 과정에서 각기 다른 이유로 "
        "부적합함이 확인되었고, 단일 측정치가 아닌 반복 검증 결과에 근거하여 두 차례 설계를 "
        "되돌렸다. 본 대회를 준비하며 오픈소스 프로젝트가 갖추어야 할 재현성과 검증 체계의 "
        "의미를 다시 생각하게 되었다.",
}

# 부록1 작성 가이드의 우선순위대로 채운다.
#   ① GPL·AGPL·LGPL 계열 → ② 핵심 기능(빠지면 안 도는 것) → ③ 프레임워크·SDK
#   → ④ 빌드·실행 도구
# 라이선스가 서로 다른 것은 한 줄로 묶지 않는다. 버전은 실제 적재된 값을 적고,
# 사용 목적에는 결합 방식을 함께 밝힌다. 팀이 직접 짠 코드는 적지 않는다.
SBOM = [
    ("1", "libgfortran", "5.0.0", "GPL-3.0-with-GCC-exception",
     "https://github.com/gcc-mirror/gcc",
     "Fortran 런타임 / NumPy 휠 번들, 동적 링크"),
    ("2", "NumPy", "1.26.4", "BSD-3-Clause", "https://github.com/numpy/numpy",
     "실행 의존성. 행렬 연산·배치 선택 / 라이브러리로 불러 씀"),
    ("3", "OpenBLAS", "0.3.23.dev", "BSD-3-Clause",
     "https://github.com/OpenMathLib/OpenBLAS",
     "행렬 연산 가속 / NumPy 휠 번들, 동적 링크"),
    ("4", "CPython", "3.11.16", "PSF-2.0", "https://github.com/python/cpython",
     "실행 런타임 / python:3.11-slim-bookworm"),
    ("5", "scikit-learn", "1.8.0", "BSD-3-Clause",
     "https://github.com/scikit-learn/scikit-learn",
     "학습 전용. TF-IDF 어휘 생성·ridge 학습 / 라이브러리로 불러 씀"),
    ("6", "SciPy", "1.17.1", "BSD-3-Clause", "https://github.com/scipy/scipy",
     "학습 전용. 희소 행렬 연산 / 라이브러리로 불러 씀"),
    ("7", "PyArrow", "23.0.1", "Apache-2.0", "https://github.com/apache/arrow",
     "자료 준비 전용. 원본 변환 / 라이브러리로 불러 씀"),
    ("8", "ossp_router (과제 제공)", "v1", "Apache-2.0",
     "https://github.com/sktelecom/ossp-2026-llm-router-challenge",
     "입출력 규격·채점 도구 / 라이브러리로 불러 씀"),
]


def set_cell(cell, text: str) -> None:
    """셀 내용을 바꾸되 양식의 문단 서식을 유지한다."""
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    para = cell.paragraphs[0]
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    lines = text.split("\n")
    run = para.add_run(lines[0])
    run.font.size = Pt(10)
    run.font.name = "맑은 고딕"
    for line in lines[1:]:
        new = cell.add_paragraph()
        new.paragraph_format.space_after = Pt(0)
        r = new.add_run(line)
        r.font.size = Pt(10)
        r.font.name = "맑은 고딕"


def main() -> int:
    doc = Document(SRC)

    info = doc.tables[2]
    row1 = list(info.rows[1].cells)
    row2 = list(info.rows[2].cells)
    set_cell(row1[1], TEAM)
    set_cell(row1[3], MEMBERS)
    set_cell(row2[1], "일반")
    set_cell(row2[3], "지정과제(SK텔레콤)")

    body = doc.tables[3]
    for row_index, text in BODY.items():
        set_cell(list(body.rows[row_index].cells)[-1], text)

    sbom = doc.tables[5]
    while len(sbom.rows) - 1 < len(SBOM):
        sbom._tbl.append(deepcopy(sbom.rows[-1]._tr))
    for i, entry in enumerate(SBOM, start=1):
        cells = list(sbom.rows[i].cells)
        for j, value in enumerate(entry):
            set_cell(cells[j], value)
    for extra in list(sbom.rows[len(SBOM) + 1:]):
        extra._tr.getparent().remove(extra._tr)

    # 붙임2 — AI 모델 활용 및 라이선스 명세
    ai = doc.tables[8]
    type_cell = list(ai.rows[1].cells)[0]
    set_cell(type_cell, type_cell.text.replace("□ 유형 3", "▣ 유형 3"))

    r3 = list(ai.rows[3].cells)
    set_cell(r3[1], "해당 없음 (기반 모델 없이 자체 학습)")
    set_cell(r3[-1], "해당 없음")

    set_cell(list(ai.rows[5].cells)[-1],
             "과제 제공 공개 Train 1,760문항 및 Dev 880문항 (총 2,640문항). "
             "프롬프트와 모델별 품질·토큰 사용량으로 구성. 출처 및 라이선스는 과제 저장소 "
             "THIRD_PARTY_NOTICES.md에 공개.")
    set_cell(list(ai.rows[6].cells)[-1],
             "과제 제공 materialize_public_data.py 절차를 그대로 사용. 개인정보가 포함되지 "
             "않은 공개 벤치마크 자료로 별도 비식별화 대상 없음. 텍스트는 소문자화 및 "
             "토큰화만 수행하고 원문 미변형.")
    set_cell(list(ai.rows[7].cells)[-1],
             f"{REPO} 의 src/routerx/artifact.npz\n"
             "(승인 절차 없이 접근 가능, 수상 시 5년간 공개 유지)")
    set_cell(list(ai.rows[8].cells)[-1],
             "파일명 : artifact.npz / 전체 가중치 배포 (ridge 회귀 7개 헤드, "
             "1.44M 파라미터) / 용량 5.5MB / 저장소 직접 커밋\n"
             "내용 : 회귀 계수 및 절편, TF-IDF 어휘·IDF(단어 60,000 · 문자 120,000), "
             "특징 정규화 통계, 등급별 안전계수 및 상한값")

    r10 = list(ai.rows[10].cells)
    set_cell(r10[1], "Apache License 2.0")
    set_cell(r10[-1], REPO + "\n(공개 유지 조건 준수)")
    set_cell(list(ai.rows[11].cells)[-1],
             "코드 작성, 실험 설계, 디버깅 보조에 Claude(Anthropic)를 활용. "
             "알고리즘 설계와 실험 결과 해석, 채택 여부의 최종 판단은 참가자가 수행.")

    # SBOM 표를 8행으로 늘리면 뒤따르는 빈 문단들이 페이지를 하나 더 밀어낸다.
    # 페이지 나누기는 남기고 내용 없는 문단만 지운다.
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    children = list(body.iterchildren())
    for index, element in enumerate(children):
        if not element.tag.endswith("}p"):
            continue
        para = Paragraph(element, doc)
        if para.text.strip():
            continue
        if "w:br" in element.xml and 'type="page"' in element.xml:
            continue          # 페이지 나누기는 유지
        prev = children[index - 1] if index else None
        if prev is not None and prev.tag.endswith("}tbl"):
            continue          # 표 직후 한 줄은 표가 붙지 않도록 남긴다
        body.remove(element)

    # 첫 쪽 작성 안내는 제출 전 삭제 대상이다. 표 인덱스가 밀리므로 마지막에 지운다.
    guide = doc.tables[0]
    guide._element.getparent().remove(guide._element)

    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"  팀명 {TEAM} / 인원 {MEMBERS} / 등록 URL 반영")
    print(f"  남은 자리표시자: 시연영상 URL")
    return 0


if __name__ == "__main__":
    sys.exit(main())
